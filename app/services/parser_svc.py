import os
import gc
import shutil
import uuid
import logging
from typing import Dict
from pathlib import Path
from fastapi import UploadFile, HTTPException
from docx import Document
import pandas as pd
import subprocess
from pathlib import Path


# 初始化日志器
logger = logging.getLogger("app.services.parser_svc")

# 确保上传目录存在
UPLOAD_DIR: Path = Path("data/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

MINERU_OUTPUT_ROOT: Path = Path("data/mineru_output")
MINERU_OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)

SUPPORTED_FORMATS = ('.pdf', '.docx', '.xlsx', '.md')
ANCHOR_COMMENT_TPL = "<!-- anchor_id:  -->"  #溯源占位符
PARAGRAPH_SEP = "\n\n"


# async：异步函数，适配FastAPI的异步特性（不阻塞其他请求）
async def save_upload_file(file: UploadFile, task_id: str) -> str:
    """
    将上传的文件持久化到磁盘，为后续 MinerU 解析做准备
    """
    #日志先行
    logger.info(f"开始保存上传文件，task_id：{task_id}，原始文件名：{file.filename}")

    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="上传文件无文件名")
        file_ext = os.path.splitext(file.filename)[1]
        file_name = f"{task_id}{file_ext}"
        dest_path = os.path.join(UPLOAD_DIR, file_name)

        # 用 task_id 重命名文件，避免多个用户上传同名文件时互相覆盖。
        suffix = Path(file.filename).suffix.lower()
        # 验证文件格式
        if suffix not in SUPPORTED_FORMATS:
            raise HTTPException(status_code=400, detail=f"不支持的文件类型: {suffix}")

        #分块复制文件内容，适合大文件
        with open(dest_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        logger.info("文件保存成功，task_id=%s，保存路径=%s", task_id, dest_path)
        return str(dest_path)

    except HTTPException:
        raise
    except Exception as e:
        logger.error("文件保存失败，task_id=%s，错误原因=%s", task_id, str(e), exc_info=True)
        raise HTTPException(status_code=500, detail=f"文件上传失败：{str(e)}")
    finally:
        await file.close()



#格式路由函数
def get_parser_engine(file_path: str) -> str:
    """根据文件后缀匹配解析引擎"""
    file_ext = os.path.splitext(file_path)[1].lower()
    engine_map = {
        '.pdf': 'mineru',
        '.docx': 'python-docx',
        '.xlsx': 'pandas',
        '.md': 'raw'
    }
    return engine_map.get(file_ext, 'raw')



# 带编码回退的文本读取函数
def _read_text_file_with_fallback(file_path: str) -> str:
    encodings = ["utf-8", "utf-8-sig", "gbk"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except Exception:
            pass

    # 如果常见编码都失败，最后用 utf-8 + ignore 强行读取。
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()



# PDF解析函数
def _parse_pdf_by_mineru(file_path: str) -> str:
    logger.info(f"使用引擎解析文件：{file_path}")

    pdf_path = Path(file_path).resolve()
    task_dir = MINERU_OUTPUT_ROOT / f"{pdf_path.stem}_{uuid.uuid4().hex}"
    task_dir.mkdir(parents=True, exist_ok=True)

    task_name = pdf_path.stem
    output_dir: Path = MINERU_OUTPUT_ROOT / task_name

    # 如果旧输出目录还在，先删除。
    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)  # 创建当前任务输出目录

    # MinerU CLI 命令
    cmd = [
        "mineru",
        "-p", str(pdf_path),
        "-o", str(output_dir),
    ]
    logger.info("调用 MinerU CLI，命令=%s", " ".join(cmd))

    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            check=True,
            timeout=3000,
        )
        if result.stdout:
            logger.info("MinerU stdout: %s", result.stdout.strip())
        if result.stderr:
            logger.warning("MinerU stderr: %s", result.stderr.strip())

    except FileNotFoundError:
        raise HTTPException(status_code=500, detail="未找到 mineru 命令，请确认 MinerU 已正确安装并加入 PATH")
    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=500, detail="MinerU 解析超时")
    except subprocess.CalledProcessError as e:
        err = e.stderr.strip() if e.stderr else str(e)
        raise HTTPException(status_code=500, detail=f"MinerU 执行失败：{err}")

    md_files = list(output_dir.rglob("*.md"))
    if not md_files:
        raise HTTPException(status_code=500, detail="MinerU 执行成功，但未生成 Markdown 文件")

    md_path = md_files[0]
    logger.info("MinerU 生成 Markdown：%s", md_path)

    try:
        content = md_path.read_text(encoding="utf-8").strip()
        if not content:
            raise HTTPException(422, "Markdown 内容为空")
        return content
    except Exception as e:
        raise HTTPException(500, f"读取 Markdown 文件失败: {str(e)}")



#DOCX解析函数
def _parse_docx(file_path: str) -> str:
    """
    解析 DOCX，段落与标题统一处理，表格按 Markdown 输出。
    """
    doc = Document(file_path)
    parts = []

    try:
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name.lower() if para.style and para.style.name else ""
            if "heading 1" in style_name:
                parts.append(f"# {text}")
            elif "heading 2" in style_name:
                parts.append(f"## {text}")
            elif "heading 3" in style_name:
                parts.append(f"### {text}")
            else:
                parts.append(text)

        for table in doc.tables:
            if not table.rows:
                continue

            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            if not any(header_cells):
                continue

            header = "| " + " | ".join(header_cells) + " |"
            sep = "| " + " | ".join(["---"] * len(header_cells)) + " |"
            rows = []

            for row in table.rows[1:]:
                row_cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(row_cells) + " |")

            table_md = "\n".join([header, sep] + rows)
            parts.append(table_md)

        content = PARAGRAPH_SEP.join(p.strip() for p in parts if p.strip()).strip()
        if not content:
            raise HTTPException(status_code=422, detail="DOCX 解析成功，但内容为空")

        return content

    finally:
        del doc


#解析Excel
def _parse_xlsx(file_path: str) -> str:
    """
    解析 Excel，每个 sheet 输出为 Markdown 表格。
    """
    excel_file = pd.ExcelFile(file_path)
    parts = []

    try:
        for sheet in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet)
            if df.empty:
                parts.append(f"### 工作表：{sheet}\n\n（空工作表）")
                continue

            df = df.fillna("")
            parts.append(f"### 工作表：{sheet}\n{df.to_markdown(index=False)}")
            del df
        content = PARAGRAPH_SEP.join(parts).strip()
        if not content:
            raise HTTPException(status_code=422, detail="Excel 解析成功，但内容为空")

        return content

    finally:
        del excel_file


#根据解析引擎提取原始文本
def parse_by_engine(engine: str, file_path: str) -> str:
    """
    根据解析引擎提取原始文本。
    """
    logger.info("使用[%s]引擎解析文件：%s", engine, file_path)

    try:
        if engine == "mineru":
            raw_content = _parse_pdf_by_mineru(file_path)
        elif engine == "python-docx":
            raw_content = _parse_docx(file_path)
        elif engine == "pandas":
            raw_content = _parse_xlsx(file_path)
        elif engine == "raw":
            raw_content = _read_text_file_with_fallback(file_path).strip()
        else:
            raise HTTPException(status_code=400, detail=f"不支持的解析引擎：{engine}")

        if not raw_content.strip():
            raise HTTPException(status_code=422, detail="解析成功但内容为空")
        logger.info("解析完成，内容长度：%s 字符", len(raw_content))
        return raw_content.strip()

    except HTTPException:
        raise
    except Exception as e:
        logger.error("[%s]引擎解析失败：%s", engine, str(e), exc_info=True)
        raise HTTPException(status_code=500, detail=f"解析失败：{str(e)}")



#文档标准化函数
def to_standard_markdown(raw_content: str) -> str:
    """
    将原始内容转为带 anchor 占位符的标准 Markdown。
    """
    logger.info("开始生成带溯源占位符的 Markdown")

    paragraphs = [p.strip() for p in raw_content.split(PARAGRAPH_SEP) if p.strip()]
    standard_paras = [f"{para} {ANCHOR_COMMENT_TPL}" for para in paragraphs]
    standard_md = PARAGRAPH_SEP.join(standard_paras)

    logger.info("标准化完成，生成 %s 个带占位符的语段", len(standard_paras))
    return standard_md


#对外主函数
def parse_document(file_path: str) -> Dict:
    """
    模块 I 主入口：解析文档并输出标准 Markdown。
    """
    logger.info("开始执行模块I解析，文件路径：%s", file_path)

    result = {"status": "success", "data": {}, "msg": "解析成功"}

    try:
        engine = get_parser_engine(file_path)
        raw_content = parse_by_engine(engine, file_path)
        standard_md = to_standard_markdown(raw_content)

        result["data"] = {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "parser_engine": engine,
            "raw_content_len": len(raw_content),
            "standard_md": standard_md,
        }
        return result
    except Exception as e:
        logger.error("模块I解析失败：%s", str(e), exc_info=True)
        result["status"] = "error"
        result["msg"] = f"解析失败：{str(e)}"
        return result
    finally:
        gc.collect()